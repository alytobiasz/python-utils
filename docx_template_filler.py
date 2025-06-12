"""
DOCX Template Filler

This script replaces bracketed fields in a Word (.docx) document with values from an Excel file.
For example, it will replace [First Name] or [First_Name] with "John" based on Excel data.

Requirements:
    - Python 3.6 or higher
    - Required Python packages:
        pip install python-docx==0.8.11    # For Word document handling
        pip install openpyxl==3.0.10       # For Excel file handling

Usage:
    1. Prepare your files:
       - DOCX template: Use bracketed fields like [Field Name] where you want replacements
       - Excel file: First row should contain headers matching the bracketed field names
         (without the brackets). For example, if [First Name] is in the document,
         the Excel should have a column header "First Name"
       - Config file: Text file with the following format:
            excel_file = path/to/data.xlsx
            template = path/to/template.docx
            output_directory = path/to/output
            filename_field1 = First Name  # Optional - uses timestamp if both fields omitted
            filename_field2 = Last Name   # Optional - uses timestamp if both fields omitted

    2. Run the script:
       python docx_template_filler.py <config_file>
    
       Example:
       python docx_template_filler.py config.txt

Note:
    - Field names in the document must match Excel headers exactly (excluding brackets)
    - Fields are case-sensitive: [First_Name] ≠ [first_name]
    - Output files will be named using the specified filename fields (or timestamp if omitted)
    - All dates are formatted as "January 1, 2025" for better readability
"""

import sys
import os
import re
import datetime as dt
import time
from docx import Document
from openpyxl import load_workbook
import traceback
from utils import format_excel_cell_date, read_config, sanitize_filename, get_unique_filename

def find_fields_in_document(doc):
    """
    Find all bracketed fields in the Word document.
    
    Args:
        doc: Word document object
        
    Returns:
        set: Set of unique field names found (without brackets)
    """
    fields = set()
    pattern = r'\[([^\]]+)\]'
    
    # Search in paragraphs
    for paragraph in doc.paragraphs:
        matches = re.finditer(pattern, paragraph.text)
        # Strip whitespace from field names
        fields.update(match.group(1).strip() for match in matches)
    
    # Search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                matches = re.finditer(pattern, cell.text)
                # Strip whitespace from field names
                fields.update(match.group(1).strip() for match in matches)
    
    return fields

def replace_fields_in_document(doc, data):
    """
    Replace all bracketed fields with corresponding values while preserving formatting.
    
    Args:
        doc: Word document object
        data (dict): Dictionary of field names and their values
    """
    # Create a mapping of field names to values (strip whitespace from keys)
    field_mapping = {}
    for key, value in data.items():
        # Handle None keys
        if key is not None:
            field_mapping[key.strip()] = value if value is not None else ''
    
    # Process paragraphs
    for paragraph in doc.paragraphs:
        replace_fields_in_paragraph(paragraph, field_mapping)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_fields_in_paragraph(paragraph, field_mapping)

def replace_fields_in_paragraph(paragraph, field_mapping):
    """
    Replace fields in a paragraph while preserving formatting.
    
    Args:
        paragraph: Paragraph object
        field_mapping (dict): Dictionary mapping field names to values
    """
    # Pattern to find field placeholders
    pattern = r'\[([^\]]+)\]'
    
    # Check if paragraph contains any fields that match our field_mapping
    # If not, we can exit early
    paragraph_has_fields = False
    for match in re.finditer(pattern, paragraph.text):
        # Get field name and strip whitespace
        field_name = match.group(1).strip()
        if field_name in field_mapping:
            paragraph_has_fields = True
            break
            
    if not paragraph_has_fields:
        return
    
    # Get all field matches in this paragraph
    field_matches = list(re.finditer(pattern, paragraph.text))
    
    # If no matches, return
    if not field_matches:
        return
    
    # Process each run, preserving formatting
    for run_index, run in enumerate(paragraph.runs):
        # Find all fields in this run
        run_matches = list(re.finditer(pattern, run.text))
        for match in run_matches:
            # Get field name and strip whitespace
            field_name = match.group(1).strip()
            if field_name in field_mapping:
                # Replace the field while preserving the run's formatting
                field_text = match.group(0)  # The full field, e.g., "[First Name]"
                value = field_mapping[field_name]
                run.text = run.text.replace(field_text, str(value))
    
    # Check if any fields were broken across runs and fix them
    # This happens when a field like [First Name] is split across multiple runs
    # For example: run1="Hello [First", run2=" Name]"
    remaining_text = paragraph.text
    merged_runs = []
    
    # Find any remaining fields that might be split across runs
    split_fields = re.finditer(pattern, remaining_text)
    for match in split_fields:
        field_text = match.group(0)  # e.g., "[First Name]"
        field_name = match.group(1).strip()  # e.g., "First Name"
        
        if field_name in field_mapping:
            # Find the starting run that contains the beginning of the field
            start_index = remaining_text.find(field_text)
            if start_index >= 0:
                # Calculate the runs that contain parts of this field
                chars_so_far = 0
                start_run = None
                end_run = None
                
                for i, run in enumerate(paragraph.runs):
                    run_length = len(run.text)
                    
                    # Check if this run contains the start of the field
                    if start_run is None and chars_so_far <= start_index < chars_so_far + run_length:
                        start_run = i
                    
                    # Check if this run contains the end of the field
                    if start_run is not None and chars_so_far + run_length >= start_index + len(field_text):
                        end_run = i
                        break
                    
                    chars_so_far += run_length
                
                # If we found both start and end, and they're different (split across runs)
                if start_run is not None and end_run is not None and start_run != end_run:
                    # Record this split field for later processing
                    merged_runs.append((start_run, end_run, field_text, str(field_mapping[field_name])))
    
    # Process any fields that were split across runs
    for start_run, end_run, field_text, replacement in reversed(merged_runs):
        # Extract the text from all affected runs
        combined_text = ''.join(paragraph.runs[i].text for i in range(start_run, end_run + 1))
        
        # Replace the field in the combined text
        new_text = combined_text.replace(field_text, replacement)
        
        # Put the text in the first run and clear others
        paragraph.runs[start_run].text = new_text
        for i in range(start_run + 1, end_run + 1):
            paragraph.runs[i].text = ''

def fill_docx_templates(config):
    """
    Fill Word document templates with data from Excel.
    
    Args:
        config (dict): Configuration dictionary containing:
            - excel_file: Path to Excel file with data
            - template: Path to Word template file
            - output_directory: Directory for output files
            - filename_field1: Optional field for filename generation
            - filename_field2: Optional field for filename generation
        
    Returns:
        tuple: (success_count, total_files) indicating number of successfully processed files
        
    Raises:
        FileNotFoundError: If the template or Excel file doesn't exist
        ValueError: If there are fields in the template not found in Excel
    """
    # Extract configuration
    excel_file = config['excel_file']
    word_template = config['template']
    output_directory = config['output_directory']
    filename_field1 = config.get('filename_field1', '')
    filename_field2 = config.get('filename_field2', '')
    
    # Create output directory if it doesn't exist
    os.makedirs(output_directory, exist_ok=True)
    
    # Verify template file exists
    if not os.path.exists(word_template):
        raise FileNotFoundError(f"Word template file not found: {word_template}")
        
    # Verify Excel file exists
    if not os.path.exists(excel_file):
        raise FileNotFoundError(f"Excel file not found: {excel_file}")
    
    # Load the template to find fields
    try:
        template_doc = Document(word_template)
    except Exception as e:
        # Convert cryptic "Package not found" error to something more meaningful
        if "Package not found" in str(e):
            raise ValueError(f"Invalid or corrupted Word document: {word_template}")
        raise
    
    template_fields = find_fields_in_document(template_doc)
    print(f"\nFound {len(template_fields)} unique fields in Word template:")
    print(", ".join(sorted(template_fields)))
    
    # Read Excel data
    wb = load_workbook(filename=excel_file, data_only=True)
    ws = wb.active
    headers = [cell.value for cell in ws[1]]
    
    # Verify all template fields exist in Excel headers
    missing_fields = []
    # Strip whitespace from all headers for consistent comparison
    stripped_headers = [h.strip() if h is not None else '' for h in headers]
    
    for field in template_fields:
        if field not in stripped_headers:
            missing_fields.append(field)
    
    if missing_fields:
        raise ValueError(f"Fields in Word template not found in Excel headers: {', '.join(missing_fields)}")
    
    # Verify filename fields exist in headers if specified
    if filename_field1:
        filename_field1 = filename_field1.strip()
        if filename_field1 not in stripped_headers:
            raise ValueError(f"Specified filename field '{filename_field1}' not found in Excel headers")
        
    if filename_field2:
        filename_field2 = filename_field2.strip()
        if filename_field2 not in stripped_headers:
            raise ValueError(f"Specified filename field '{filename_field2}' not found in Excel headers")
    
    # Create a mapping from stripped headers to original headers
    header_mapping = {h.strip() if h is not None else '': h for h in headers}
    
    if not filename_field1 and not filename_field2:
        print("No filename fields specified - using timestamps for output files")
    
    # Count total non-empty rows
    total_files = sum(1 for row in ws.iter_rows(min_row=2) if any(cell.value for cell in row))
    processed_count = 0
    success_count = 0
    
    # Process each row
    for row_cells in ws.iter_rows(min_row=2):
        row = [cell.value for cell in row_cells]
        if not any(row):  # Skip empty rows
            continue
        
        processed_count += 1
        start_time = time.time()
        
        try:
            # Create data dictionary
            data = {}
            for i, cell in enumerate(row_cells):
                if i < len(headers):
                    header = headers[i]
                    if header is not None:
                        data[header.strip()] = format_excel_cell_date(cell)
            
            # Generate output filename from specified fields
            if filename_field1 or filename_field2:
                # Use empty string if the value is None or not in the data dictionary
                field1_value = str(data.get(filename_field1) or '').strip()
                field2_value = str(data.get(filename_field2) or '').strip()
                filename = f"{field1_value} {field2_value}".strip()
            else:
                # Use timestamp if no fields specified
                filename = dt.datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # Sanitize filename
            filename = sanitize_filename(filename)
            
            # Create output path and handle duplicates
            base_path = os.path.join(output_directory, filename)
            docx_path = get_unique_filename(base_path, "docx")
            
            # Create and save the filled document
            doc = Document(word_template)
            replace_fields_in_document(doc, data)
            doc.save(docx_path)
            
            success_count += 1
            elapsed_time = time.time() - start_time
            print(f"Processed {processed_count}/{total_files}: {os.path.basename(docx_path)} in {elapsed_time:.1f} seconds")
            
        except Exception as e:
            # Log the error
            print(f"Error processing row {processed_count}: {str(e)}")
            print("Stack trace:")
            traceback.print_exc()
            # Re-raise the exception to propagate it
            raise
    
    # Print summary
    print("\nProcessing Summary:")
    print(f"Total files processed: {success_count}/{total_files}")
    print(f"Output directory: {os.path.abspath(output_directory)}")
    
    return success_count, total_files

def main():
    """Main function to handle command line arguments."""
    if len(sys.argv) != 2:
        print("Usage: python docx_template_filler.py <config_file>")
        sys.exit(1)
    
    config = read_config(sys.argv[1])
    fill_docx_templates(config)

if __name__ == "__main__":
    main() 